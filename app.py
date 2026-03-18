import streamlit as st
import pandas as pd
import os
import datetime
import re
import io
import time
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from crewai import Agent, Task, Crew, LLM
from crewai_tools import SerperDevTool

# ============================================================================
# 1. 환경 및 보안 설정
# ============================================================================
st.set_page_config(page_title="구매지원팀 통합 분석 시스템 (Gemini)", layout="wide")

if "GEMINI_API_KEY" in st.secrets:
    api_key = st.secrets["GEMINI_API_KEY"]
    os.environ["GEMINI_API_KEY"] = api_key
    os.environ["GOOGLE_API_KEY"] = api_key 
    os.environ["SERPER_API_KEY"] = st.secrets.get("SERPER_API_KEY", "")
else:
    st.error("⚠️ Streamlit Cloud의 Secrets 설정을 완료해 주세요.")

url = "https://docs.google.com/spreadsheets/d/e/2PACX-1vST3eDNhF1GLc231d4RdAnSCb8DnSznnZ4lJfPxxmtIHIcuEXbvFmrBI9LRdbURog-ik09vSOHTOAMp/pub?output=csv"

@st.cache_data(ttl=600)
def load_data():
    try:
        data = pd.read_csv(url)
        # 특정 품목명을 가독성 좋게 변경
        data['품목'] = data['품목'].replace('가수분해소고기농축물(호주)', '호주산 쇠고기 (가수분해농축물)')
        data['날짜'] = pd.to_datetime(data['날짜'])
        return data.sort_values(['품목', '날짜'])
    except Exception as e:
        st.error(f"데이터 로드 중 오류 발생: {e}")
        return None

def markdown_to_docx_stream(markdown_text):
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
    lines = markdown_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        if line.startswith('# '): doc.add_heading(line[2:], level=0).alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif line.startswith('## '): doc.add_heading(line[3:], level=1)
        elif line.startswith('### '): doc.add_heading(line[4:], level=2)
        else: doc.add_paragraph(line)
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

df_raw = load_data()

# ============================================================================
# 2. 통계 계산 및 핵심 이슈 품목 추출 (주/월/연 각 2개씩 선정)
# ============================================================================
def calculate_all_stats(df):
    df = df.copy()
    df['연주'] = df['날짜'].dt.to_period('W').astype(str)
    df['연월'] = df['날짜'].dt.to_period('M').astype(str)
    df['연도'] = df['날짜'].dt.year
    def get_stats(df_group, col_name):
        grouped = df_group.groupby(['품목', '단위', col_name])['y'].mean().reset_index()
        grouped.columns = ['품목', '단위', '기간', '평균시세']
        grouped['이전시세'] = grouped.groupby('품목')['평균시세'].shift(1)
        grouped['증감률'] = ((grouped['평균시세'] - grouped['이전시세']) / grouped['이전시세'] * 100).round(2)
        return grouped.fillna(0)
    return get_stats(df, '연주'), get_stats(df, '연월'), get_stats(df, '연도')

def get_critical_items(w_df, m_df, y_df):
    w_top = w_df[w_df['기간'] == w_df['기간'].max()].nlargest(1, '증감률')['품목'].tolist()
    w_bot = w_df[w_df['기간'] == w_df['기간'].max()].nsmallest(1, '증감률')['품목'].tolist()
    m_top = m_df[m_df['기간'] == m_df['기간'].max()].nlargest(1, '증감률')['품목'].tolist()
    m_bot = m_df[m_df['기간'] == m_df['기간'].max()].nsmallest(1, '증감률')['품목'].tolist()
    y_top = y_df[y_df['기간'] == y_df['기간'].max()].nlargest(1, '증감률')['품목'].tolist()
    y_bot = y_df[y_df['기간'] == y_df['기간'].max()].nsmallest(1, '증감률')['품목'].tolist()
    return list(set(w_top + w_bot + m_top + m_bot + y_top + y_bot))

# ============================================================================
# 3. 메인 대시보드 (기존 3x3 레이아웃)
# ============================================================================
if df_raw is not None:
    weekly_df, monthly_df, yearly_df = calculate_all_stats(df_raw)
    st.title("📊 원자재 시세 실시간 분석 및 전문 AI 보고서")
    st.info(f"데이터 업데이트 시각: {datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    display_cols = ['품목', '평균시세', '단위', '증감률']
    def format_df(df): return df[display_cols].style.format({'평균시세': '{:,.2f}', '증감률': '{:+.2f}%'})

    for title, func in [("🔍 기간별 전체 시세 현황", None), ("📈 가격 상승 TOP 5", "nlargest"), ("📉 가격 하락 TOP 5", "nsmallest")]:
        st.header(title)
        c1, c2, c3 = st.columns(3)
        for col, data, p_name in zip([c1, c2, c3], [weekly_df, monthly_df, yearly_df], ["🗓️ 주간", "📅 월간", "📂 연간"]):
            with col:
                st.subheader(p_name)
                curr = data[data['기간'] == data['기간'].max()]
                disp = getattr(curr, func)(5, '증감률') if func else curr
                st.dataframe(format_df(disp), use_container_width=True, hide_index=True)
        st.divider()

    # ============================================================================
    # 4. 전문 AI 분석 섹션 (1년치 분기별 예측 강화)
    # ============================================================================
    st.header("🔮 1년 전망 단가 예측 보고서 (Gemini 2.5 Flash)")
    critical_items = get_critical_items(weekly_df, monthly_df, yearly_df)
    
    today_str = datetime.date.today().strftime('%Y년 %m월 %d일')
    future_limit = (datetime.date.today() + datetime.timedelta(days=365)).strftime('%Y년 %m월')

    st.write(f"🔔 **AI 예측 대상:** {', '.join(critical_items)}")
    st.caption(f"※ {today_str} 기준 정보를 바탕으로 **{future_limit}까지의 향후 1년 시세**를 분기별로 예측합니다.")

    if st.button("🚀 1년 단가 예측 시작"):
        if not os.environ.get("GEMINI_API_KEY"):
            st.error("🚨 API 키가 설정되지 않았습니다.")
        else:
            search_tool = SerperDevTool()
            gemini_llm = LLM(model="gemini/gemini-2.5-flash", api_key=os.environ["GEMINI_API_KEY"])

            with st.status("향후 1개년 시장 시나리오 분석 중...", expanded=True) as status:
                analyst = Agent(
                    role="장기 시장 수급 예측가", 
                    goal=f"오늘({today_str}) 이후의 데이터를 분석하여 향후 1년(4개 분기)의 단가 흐름을 예측", 
                    backstory="당신은 글로벌 수급 트렌드를 읽어내는 베테랑 분석가입니다. 단기적인 변동보다 향후 1년간의 분기별 단가 추이를 예측하는 데 특화되어 있습니다.", 
                    llm=gemini_llm, 
                    tools=[search_tool]
                )
                procurement = Agent(
                    role="전략적 연간 구매 설계자", 
                    goal="예측된 1년치 단가 추이에 따른 분기별 최적 구매 로드맵 제안", 
                    backstory="당신은 1년치 예산을 효율적으로 집행하기 위해 어느 분기에 대량 매수를 하고 어느 분기에 관망해야 할지를 결정하는 구매 전략의 대가입니다.", 
                    llm=gemini_llm
                )

                all_reports = []
                progress_bar = st.progress(0)
                
                for idx, item in enumerate(critical_items):
                    st.write(f"🔮 **{item}** 1년 전망 분석 중... ({idx+1}/{len(critical_items)})")
                    
                    t1 = Task(
                        description=f"""
                        품목: {item}
                        현재 날짜: {today_str}
                        미션: 
                        1. 오늘 이후의 최신 뉴스 및 시장 동향을 검색하여 {item}의 향후 1년(4개 분기) 시세를 분석하세요.
                        2. **[1년 분기별 시나리오]** 섹션을 만들어 각 분기별(예: 2026년 Q2, Q3, Q4, 2027년 Q1)로 단가 흐름을 구체적으로 예측하세요.
                        3. '앞으로 일어날 주요 이벤트'가 각 분기 가격에 어떤 영향을 줄지 설명하세요.
                        """, 
                        expected_output=f"{item}의 향후 1년 분기별 단가 변동 시나리오 보고서", 
                        agent=analyst
                    )
                    
                    t2 = Task(
                        description=f"""
                        위의 {item} 1년 예측 결과를 바탕으로 연간 구매 전략을 수립하세요.
                        - **분기별 구매 로드맵**: 각 분기별로 '공격적 매수', '안정적 확보', '구매 대기' 중 하나를 선택하고 이유를 적으세요.
                        - **최적 매수 타이밍**: 1년 중 단가가 가장 저렴할 것으로 예상되는 '골든 타임'을 콕 짚어주세요.
                        - **리스크 변수**: 향후 1년간 주의해야 할 외부 변수를 지정하세요.
                        """, 
                        expected_output=f"{item}의 향후 1년 구매 실행 로드맵", 
                        agent=procurement
                    )
                    
                    crew = Crew(agents=[analyst, procurement], tasks=[t1, t2], max_rpm=1)

                    success = False
                    for attempt in range(3):
                        try:
                            report_out = crew.kickoff()
                            all_reports.append(report_out.raw)
                            success = True
                            break
                        except Exception as e:
                            if "429" in str(e):
                                time.sleep(15 * (attempt + 1))
                            else:
                                break
                    
                    if not success:
                        all_reports.append(f"### {item}\n분석 한도 초과로 1년 리포트 생성에 실패했습니다.")
                    
                    time.sleep(7)
                    progress_bar.progress((idx + 1) / len(critical_items))

                # [중요] 보고서 통합 및 취소선(~~) 제거 처리
                final_report_md = f"# 📑 [연간전략] 1년 단가 예측 및 구매 로드맵 ({today_str} 발행)\n\n" + "\n\n---\n\n".join(all_reports)
                
                # 마크다운 취소선 문법인 '~~'를 아예 제거하여 텍스트가 지워지는 현상 방지
                final_report_md = final_report_md.replace('~~', '')
                
                status.update(label="✅ 모든 핵심 품목 1년 예측 완료!", state="complete", expanded=False)

            st.markdown(final_report_md)
            docx_file = markdown_to_docx_stream(final_report_md)
            st.download_button(label="📄 1년 예측 보고서 다운로드 (Word)", data=docx_file, file_name=f"Annual_Prediction_Report_{datetime.date.today()}.docx")
